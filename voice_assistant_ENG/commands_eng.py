from get_audio import get_audio_eng
from speak import speak_eng
import os
from time import strftime, sleep
import docx
import pyautogui as pg
import winshell
import matplotlib.pyplot as plt
import requests
import pyodbc
from voice_assistant_ENG import main_eng
import pymongo


def timer(mongo_collect, user):
    speak_eng("OK, I turn on timer")
    speak_eng("What I will remind you?")
    reminder_name = get_audio_eng().lower()
    speak_eng("How many minutes?")
    flag = 0
    while flag == 0:
        digit = get_audio_eng()
        for char in digit:
            if char.isdigit():
                digit = char
                break
        try:
            minutes = int(digit)
            minutes = minutes * 60
            flag += 1
            sleep(minutes)
            outcome = speak_eng(f"You should {reminder_name}")
            mongo_collect.insert_one(
                {"user": user, "action": "Turn on the timer",
                 "details": [
                     {"[ACTION]": "TIMER ON"},
                     {"[TIME]": f'{strftime("%c")}'},
                     {"[REMIND ABOUT]": f'{reminder_name}'},
                     {"[HOW MANY MINUTES]": f'{minutes}'}
                 ],
                 "status": "SUCCESS",
                 "outcome": f"{outcome }"
                 }
            )
        except ValueError:
            outcome = speak_eng("Please, say some digit")
            mongo_collect.insert_one(
                {"user": user, "action": "Turn on the timer",
                 "details": [
                     {"[ACTION]": "TIMER ON"},
                     {"[TIME]": f'{strftime("%c")}'},
                     {"[REMIND ABOUT]": f"{reminder_name}"},
                     {"[HOW MANY MINUTES]": f"{minutes}"}
                 ],
                 "status": "FAIL",
                 "outcome": f"{outcome}"
                 }
            )


def time(mongo_collect, user):
    outcome = speak_eng(f'Now is {strftime("%X")}')
    mongo_collect.insert_one(
        {"user": user, "action": "Voicing current time",
         "details": [
             {"[ACTION]": "CURRENT TIME"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f"{outcome}"
         }
    )


def money(mongo_collect, user):
    city = 'Гомель'
    api_not_json = requests.get(f'https://belarusbank.by/api/kursExchange?city={city}')
    api = api_not_json.json()
    speak_eng("What kind of currency do you want to get?")
    currency = get_audio_eng()
    speak_eng("Do you want to sell it or to buy?")
    transaction = get_audio_eng()
    if 'dollar' or '$' in currency:
        currency = 'USD'
    elif 'euro' in currency:
        currency = 'EUR'
    elif 'ruble' in currency:
        currency = 'RUB'
    if 'sell' in transaction:
        value = api[0][f'{currency}_in']
        outcome = speak_eng(f"You can sell {currency} for {value}")
    else:
        value = api[0][f'{currency}_out']
        outcome = speak_eng(f"You can buy {currency} for {value}")
    mongo_collect.insert_one(
        {"user": user, "action": "Getting info about currency",
         "details": [
             {"[ACTION]": "CURRENCY\'s INFO "},
             {"[TIME]": f'{strftime("%c")}'},
             {"[CURRENCY]": f'{currency}'},
             {"[TRANSACTION]": f'{transaction}'}
         ],
         "status": "SUCCESS",
         "outcome": f"{outcome}"
         }
    )


def weather(mongo_collect, user):
    appid = "66eb8cf1abebae782e1d3700a75e203e"

    def form_url_string(s_request):
        nonlocal appid
        s_appid = "&APPID=" + appid
        s_template = "http://api.openweathermap.org/data/2.5/" + s_request + s_appid
        return s_template

    s_city = "Gomel"
    s_country = "BY"
    s_request = f"weather?q={s_city},{s_country}&units=metric"
    s_search_url = form_url_string(s_request)
    try:
        res = requests.get(s_search_url)
        data = res.json()
        outcome = speak_eng(
            f"Weather now:\n Conditions: {data['weather'][0]['description']}\n Temperature: {data['main']['temp']}")
    except Exception as e:
        speak_eng("Sorry I couldn't find your city")
    mongo_collect.insert_one(
        {"user": user, "action": "Voicing the current weather",
         "details": [
             {"[ACTION]": "THE CURRENT WEATHER"},
             {"[TIME]": f'{strftime("%c")}'},
             {"[CITY]": f'{s_city}'},
             {"[COUNTRY]": f'{s_country}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def creating_word_file(mongo_collect, user):
    speak_eng('Say a file name')
    file_name = get_audio_eng()
    document = docx.Document()
    speak_eng(f'Your files name : {file_name}')
    speak_eng('Do you want to create an empty file?')
    accept = get_audio_eng()
    list_words = ('yes', 'yeah', 'create an empty file')
    for i in list_words:
        if i in accept:
            document.save(f'{file_name}.docx')
            os.system(f'{file_name}.docx')
            break
    else:
        speak_eng('Say some text which will be in the title')
        title_word = get_audio_eng()
        first_letter = title_word[0].upper()
        document.add_heading(f'{first_letter}{title_word[1:]}', 0)
        speak_eng('Say some text which will be contain your file')
        just_text = get_audio_eng()
        first_letter = just_text[0].upper()
        p = document.add_paragraph(f'{first_letter}{just_text[1:]}  ')
        p.add_run(f'{first_letter}{just_text[1:]}  ').bold = True
        p.add_run(f'{first_letter}{just_text[1:]}.').italic = True
        document.save(f'{file_name}.docx')
        os.system(f'{file_name}.docx')
    outcome = speak_eng("The file was created")
    mongo_collect.insert_one(
        {"user": user, "action": "Creating the word file",
         "details": [
             {"[ACTION]": "CREATE THE WORD FILE"},
             {"[TIME]": f'{strftime("%c")}'},
             {"[FILE's NAME]": f'{file_name}'},
             {"[FILE's TYPE]": f'{""}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def open_my_computer(mongo_collect, user):
    os.startfile(r"shell:mycomputerfolder")
    outcome = speak_eng('I opened my computer')
    mongo_collect.insert_one(
        {"user": user, "action": "Opening the folder My Computer",
         "details": [
             {"[ACTION]": "OPEN MY COMPUTER"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def cleaning_recycle_bin(mongo_collect, user):
    speak_eng("Do you really want it?")
    text = get_audio_eng()
    if "yes".lower() in text:
        outcome = speak_eng("I'm cleaning")
        try:
            winshell.recycle_bin().empty(confirm=False, show_progress=False, sound=False)
        except BaseException:
            outcome = speak_eng("Sorry, I can't do this because Your recycle bin is empty")
    elif "no".lower() in text:
        outcome = speak_eng("Good, I didn't clean")
    mongo_collect.insert_one(
        {"user": user, "action": "Cleaning the recycle bin",
         "details": [
             {"[ACTION]": "CLEAN THE RECYCLE BIN"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def opening_the_browser(mongo_collect, user):
    pg.hotkey('win', 'r')
    speak_eng("I'm opening the browser")
    pg.typewrite('msedge')
    pg.press('enter')
    speak_eng("What do you want to find?")
    text = get_audio_eng()
    speak_eng(f"I'm searching : {text}")
    pg.typewrite(text)
    print(text)
    pg.typewrite(["enter"])
    outcome = speak_eng(f"I opened the browser and the searching text is {text}")
    mongo_collect.insert_one(
        {"user": user, "action": "Opening the browser",
         "details": [
             {"[ACTION]": "OPEN THE BROWSER"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def check_command(mongo_collect, user):
    speak_eng("What kind of command do you want to check?")
    name_command = get_audio_eng()
    speak_eng("How many times?")
    count = 0
    while count == 0:
        digit = get_audio_eng()
        for char in digit:
            if char.isdigit():
                digit = char
                break
        try:
            kol = int(digit)
            count = 1
        except ValueError:
            speak_eng("Please, say some digit")
    x = []
    y = []
    for number in range(kol):
        if number != 0:
            speak_eng("Please, say the name of command")
            name_command = get_audio_eng()
        main_eng.main_eng_va(name_command)
        speak_eng('How outcome for you? Please say 1 if outcome is good else enter 0')
        x.append(number)
        y.append(int(input()))
    fig = plt.figure()
    ax = fig.add_subplot(111)
    speak_eng("Please, enter the name of your figure")
    name = input()
    ax.set(title=f'{name}')
    ax.set_xlabel('Number of function calls')
    ax.set_ylabel('Score for the outcome')
    plt.plot(x, y)
    plt.show()
    outcome = speak_eng("I checked commands. Please look at the screen. I built a figure")
    mongo_collect.insert_one(
        {"user": user, "action": "Checking commands",
         "details": [
             {"[ACTION]": "CHECK THE COMMANDS"},
             {"[TIME]": f'{strftime("%c")}'},
             {"[NAME OF COMMAND]": f'{name_command}'},
             {"[HOW MANY TIMES CHECK]": f'{kol}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def calculator(mongo_collect, user):
    speak_eng("What do you want to calculate?")
    text = get_audio_eng()
    text = text.split()
    try:
        if "+" in text[1] or "plus" in text[1]:
            outcome = speak_eng(f'{text[0]} plus {text[2]} is {int(text[0]) + int(text[2])}')
        elif "-" in text[1] or "min" in text[1]:
            outcome = speak_eng(f'{text[0]} minus {text[2]} is {int(text[0]) - int(text[2])}')
        elif "*" in text[1] or "multiply" in text[1] or "multiplied" in text[1] or 'х' in text[1]:
            outcome = speak_eng(f'{text[0]} multiply {text[2]} is {int(text[0]) * int(text[2])}')
        elif "/" in text[1] or "divide" in text[1] or "divided" in text[1]:
            outcome = speak_eng(f'{text[0]} divide {text[2]} is {int(text[0]) / int(text[2])}')
    except ZeroDivisionError:
        outcome = speak_eng("I can't to divide on 0")
    except Exception:
        outcome = speak_eng("Sorry, happens some mistake")
    mongo_collect.insert_one(
        {"user": user, "action": "Calculating",
         "details": [
             {"[ACTION]": "CALCULATE THE NUMBERS"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def capability(mongo_collect, user):
    conn = pyodbc.connect('Driver={SQL Server};'
                          'Server=USER-PC\SQLEXP;'
                          'Database=Voice_assistant;'
                          'Trusted_Connection=yes;')
    cursor = conn.cursor()
    try:
        cursor.execute('''drop table dbo.SKILLS_EN''')
    except pyodbc.ProgrammingError:
        pass
    cursor.execute('''Create table dbo.SKILLS_EN(
                                        Name varchar(60) unique
                                      )'''
                   )
    cursor.execute('''Insert into dbo.SKILLS_EN values
                                        ('Turn the timer on'),
                                        ('Speak a Current time'),
                                        ('Speak a Current weather'),
                                        ('Create Word-file'),
                                        ('Open my computer'),
                                        ('Clean the recycle bin'),
                                        ('Open Chrome'),
                                        ('Calculate simple mathematical actions')
                                    ''')
    cursor.execute('SELECT Name FROM Voice_assistant.dbo.SKILLS_EN order by Name')
    outcome = []
    outcome.append(speak_eng("I can: "))
    for row in cursor:
        outcome.append(speak_eng(row[0]))
    outcome = '\n'.join(outcome)
    conn.commit()
    conn.close()
    mongo_collect.insert_one(
        {"user": user, "action": "Voicing the Capability",
         "details": [
             {"[ACTION]": "CAPABILITIES OF VA"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def search(mongo_collect, user):
    pg.hotkey('winleft', 's')
    speak_eng("I'm ready for searching. What's name of the file?")
    file_name = get_audio_eng()
    pg.typewrite(f'{file_name}')
    outcome = speak_eng("Look on the screen, I finished the searching")
    mongo_collect.insert_one(
        {"user": user, "action": "Searching",
         "details": [
             {"[ACTION]": "SEARCH THE FILE IN THE PS"},
             {"[TIME]": f'{strftime("%c")}'},
             {"[NAME OF SEARCHING THING]": f'{file_name}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def change_keyboard_lang(mongo_collect, user):
    pg.hotkey('win', 'space')
    outcome = speak_eng('I changed the language')
    mongo_collect.insert_one(
        {"user": user, "action": "Changing keyboard language",
         "details": [
             {"[ACTION]": "ChANGE KEYBOARD LANGUAGE"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def volume(action, mongo_collect, user):
    if action == "down":
        pg.press('volumedown')
    else:
        pg.press('volumeup')
    outcome = speak_eng("I changed volume")
    mongo_collect.insert_one(
        {"user": user, "action": "Changing the volume",
         "details": [
             {"[ACTION]": "CHANGE THE VOLUME"},
             {"[TIME]": f'{strftime("%c")}'},
             {"[TYPE OF CHANGING]": f'CHANGING {action}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def task_manager(mongo_collect, user):
    outcome = speak_eng("I'm opening task manager")
    pg.hotkey('ctrl', 'shift', 'esc')
    mongo_collect.insert_one(
        {"user": user, "action": "Opening task manager",
         "details": [
             {"[ACTION]": "OPEN TASK MANAGER"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f'{outcome}'
         }
    )


def history(name, cursor, mongo_collect):
    speak_eng(f"Do you want to see your history? Please say either yes or no")
    answer = get_audio_eng()
    if "yes" in answer:
        cursor.execute(f"SELECT COMMAND_N FROM Voice_assistant.dbo.HISTORY where Name = '{name}'")
    else:
        speak_eng("Please say the full name of user who you want to check")
        another_name = get_audio_eng()
        cursor.execute(f"SELECT COMMAND_N FROM Voice_assistant.dbo.HISTORY where NAME = '{another_name}'")
        speak_eng(f"The following actions was taken by {name}")
    outcome = []
    outcome.append(speak_eng(f"The history of {name} are: "))
    for row in cursor:
        outcome.append(row[0])
    outcome = '\n '.join(outcome)
    print(type(outcome))
    print(outcome)
    mongo_collect.insert_one(
        {"user": name, "action": "Showing history",
         "details": [
             {"[ACTION]": "SHOW THE HISTORY"},
             {"[TIME]": f'{strftime("%c")}'}
         ],
         "status": "SUCCESS",
         "outcome": f"{outcome}"
         }
    )



